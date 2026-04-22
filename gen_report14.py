# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(3); s.right_margin = Cm(1.5)

# ── helpers ──────────────────────────────────────────────────
def sf(run, size=14, bold=False, italic=False, mono=False):
    run.font.name = 'Courier New' if mono else 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def pc(text, size=14, bold=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(21)
    sf(p.add_run(text), size=size, bold=bold)

def pr(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = Pt(21)
    p.paragraph_format.space_after  = Pt(4)
    sf(p.add_run(text), size=size)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(21)
    sf(p.add_run(text.upper()), size=14, bold=True)

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(21)
    sf(p.add_run(text), size=14, bold=True)

def h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(21)
    sf(p.add_run(text), size=14, bold=True, italic=True)

def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.line_spacing      = Pt(21)
    sf(p.add_run(text), size=14)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(21)
    p.paragraph_format.space_after  = Pt(2)
    sf(p.add_run(text), size=14)

def code(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(17)
    sf(p.add_run(text), size=11, mono=True)

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(2)
        pp.paragraph_format.space_after  = Pt(2)
        sf(pp.add_run(h), size=11, bold=True)
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.space_before = Pt(1)
            pp.paragraph_format.space_after  = Pt(1)
            sf(pp.add_run(str(ct)), size=11)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
pc("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ", 14, bold=True, sb=0, sa=2)
pc("КИЇВСЬКИЙ НАЦІОНАЛЬНИЙ ЕКОНОМІЧНИЙ УНІВЕРСИТЕТ", 14, bold=True, sa=2)
pc("імені Вадима Гетьмана", 14, sa=2)
pc("Факультет комп'ютерних наук та інформаційних технологій (ФКІСІТ)", 14, sa=36)
pc("ЗВІТ", 16, bold=True, sa=4)
pc("з лабораторної роботи №14", 14, bold=True, sa=4)
pc("на тему:", 14, sa=4)
pc("«Підготовка представлень сторінок", 14, bold=True, sa=2)
pc("для менеджера web-сайту»", 14, bold=True, sa=36)
pr("Виконав: студент групи 371\nЗівак Сергій\n\nПеревірив: Юрій Лозовик")
pc("", sb=30)
pc("Київ — 2026", 14, sb=20, sa=0)
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# ТЕМА ТА МЕТА
# ════════════════════════════════════════════════════════════
h1("Тема та Мета лабораторної роботи")
body("Тема: Підготовка представлень сторінок для менеджера web-сайту.")
body("Мета: розробити та дослідити FreeMarker-шаблони представлень (View) для "
     "менеджерської панелі веб-застосунку PayFlow. Вивчити принципи проектування "
     "інтерфейсів з обмеженим доступом відповідно до ролі MANAGER, реалізувати "
     "візуальне розмежування повноважень між адміністраторською та менеджерською "
     "панелями, а також проаналізувати механізм захисту представлень через "
     "HttpSession-перевірку у ManagerController.")
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. ТЕОРЕТИЧНІ ВІДОМОСТІ
# ════════════════════════════════════════════════════════════
h1("1. Теоретичні відомості")

h2("1.1. Роль менеджера у веб-застосунках (принцип найменших привілеїв)")
body("Принцип найменших привілеїв (Principle of Least Privilege, PoLP) — один з "
     "фундаментальних принципів інформаційної безпеки. Відповідно до нього, "
     "кожен суб'єкт системи (користувач, процес, роль) повинен мати лише ті "
     "права доступу, які мінімально необхідні для виконання його функцій. "
     "Будь-яке розширення привілеїв понад необхідний мінімум збільшує потенційну "
     "поверхню атаки та ризик зловживань.")
body("У контексті веб-застосунків роль менеджера зазвичай займає проміжне місце "
     "між клієнтом (CLIENT) та адміністратором (ADMIN). Менеджер отримує доступ "
     "до агрегованих аналітичних даних, але не може вносити зміни у систему: "
     "не управляє обліковими записами, не редагує транзакції, не може призначати "
     "ролі. Таке розмежування дозволяє залучати менеджерський персонал до "
     "моніторингу бізнес-показників без ризику несанкціонованих змін.")
body("У проєкті PayFlow роль MANAGER визначена в enum UserRole і відповідає "
     "наступному набору привілеїв: перегляд статистики та метрик (dashboard), "
     "перегляд журналу транзакцій з фільтрацією та пошуком, перегляд сторінки "
     "звітів (наразі — заглушка). Управління користувачами, підтвердження "
     "транзакцій та зміна налаштувань системи недоступні.")

h2("1.2. Особливості проектування представлень для обмеженого доступу")
body("При проектуванні інтерфейсів з обмеженим доступом дотримуються кількох "
     "ключових підходів:")
for item in [
    "Відсутність елементів керування у View: кнопки, форми та посилання на заборонені дії "
    "фізично відсутні у шаблоні — не приховані CSS/JS, а саме відсутні. "
    "Це унеможливлює їх активацію навіть при маніпуляціях з DOM.",
    "Явна комунікація обмежень: readonly-banner і readonly-notice у менеджерських шаблонах "
    "чітко повідомляють користувачу про режим перегляду, не залишаючи простору "
    "для непорозуміння щодо доступних можливостей.",
    "Спрощена навігація: навігаційне меню менеджера містить лише ті пункти, "
    "до яких є доступ (3 пункти замість 5 у адміністратора).",
    "Візуальна диференціація: різна колірна схема (смарагдова #064e3b для менеджера "
    "проти індиго #1e1b4b для адміністратора) дозволяє миттєво розрізнити "
    "відкриту сесію навіть без читання текстових підказок.",
    "Захист на рівні сервера (Defense in Depth): навіть якщо користувач "
    "вручну введе URL /admin/users, відсутність сесійного атрибута adminUser "
    "призведе до перенаправлення, а не до відображення сторінки.",
]:
    bullet(item)

h2("1.3. FreeMarker директиви що використовуються у шаблонах менеджера")
body("Шаблони менеджерської панелі активно застосовують наступні FreeMarker-конструкції:")
tbl(
    ["Директива / Конструкція", "Приклад використання", "Призначення"],
    [
        ["${variable}",            "${manager.firstName}",                   "Виведення значення змінної"],
        ["${var!\"default\"}",     "${p.transactionId!\"—\"}",               "Виведення або default якщо null"],
        ["${var?string[\"0.00\"]}", "${totalBalance?string[\"0.00\"]}",      "Форматування числа з 2 знаками"],
        ["${var?substring(0,1)}",  "${manager.firstName?substring(0,1)}",    "Перший символ рядка (ініціали)"],
        ["${var?c}",               "${p.createdAt.year?c}",                  "Число → рядок без форматування"],
        ["${var?size}",            "${topUsers?size}",                       "Розмір колекції"],
        ["${var?string[\"0\"]}",   "${successRate?string[\"0\"]}%",          "Форматування без дробової частини"],
        ["variable??",             "<#if manager??>",                        "Перевірка що змінна не null"],
        ["<#if c>…</#if>",         "<#if (pendingApprovals > 0)>",           "Умовний блок"],
        ["<#if c>…<#else>…</#if>", "<#if manager??>…<#else>M</#if>",        "Умова з гілкою else"],
        ["<#if c>…<#elseif c>…",   "порівняння типу/статусу транзакції",     "Багатогілкова умова"],
        ["<#list col as item>",    "<#list topUsers as u>",                  "Ітерація по колекції"],
        ["<#list a..b as i>",      "<#list 0..(totalPages-1) as i>",         "Цикл по числовому діапазону"],
        ["item_index",             "<#assign rank = u_index + 1>",           "Індекс елемента (0-based)"],
        ["item_has_next",          "<#if u_has_next>border<#else>none</#if>","Чи є наступний елемент"],
        ["<#assign name = val>",   "<#assign successRate = 0>",              "Локальна змінна у шаблоні"],
        ["<#sep>,</#sep>",         "<#list chartData as d>${d}<#sep>,</#sep>","Роздільник між елементами"],
        ["<#-- comment -->",       "<#-- @ftlvariable name=\"manager\"... -->","Коментар (не виводиться в HTML)"],
    ],
    [4.5, 5.5, 7.0]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. ЗАГАЛЬНА СТРУКТУРА МЕНЕДЖЕРСЬКОЇ ПАНЕЛІ
# ════════════════════════════════════════════════════════════
h1("2. Загальна структура менеджерської панелі")

h2("2.1. Таблиця шаблонів менеджерської панелі")
tbl(
    ["Шаблон", "Маршрут (URL)", "Призначення", "Ключова відмінність від admin/"],
    [
        ["manager/login.ftl",
         "GET /manager/login",
         "Форма входу менеджера",
         "Зелена колірна гама (#064e3b), іконка bi-speedometer2, POST на /manager/login"],
        ["manager/dashboard.ftl",
         "GET /manager/dashboard",
         "Аналітичний дашборд (тільки перегляд)",
         "readonly-banner замість pending-banner, зелені кольори Chart.js та акцентів"],
        ["manager/transactions.ftl",
         "GET /manager/transactions",
         "Перегляд журналу транзакцій",
         "Відсутня колонка Actions (без approve/reject форм), readonly-notice badge"],
        ["manager/reports.ftl",
         "GET /manager/reports",
         "Заглушка сторінки звітів",
         "Зелена іконка/кнопка, посилання на /manager/dashboard"],
    ],
    [4.0, 3.8, 4.5, 4.7]
)

h2("2.2. Спільні елементи всіх шаблонів менеджерської панелі")
body("Кожен з чотирьох шаблонів (крім login.ftl) містить однакову структуру: "
     "фіксований лівий сайдбар (220px), sticky топбар та контентну область. "
     "Ці елементи повторюються у кожному шаблоні (без механізму include/extends, "
     "що є технічним боргом проєкту).")

h3("Код сайдбара з manager/dashboard.ftl")
code("<aside class=\"mgr-sidebar\">")
code("  <div class=\"sidebar-brand\">")
code("    <div class=\"brand-icon\">")
code("      <i class=\"bi bi-speedometer2\"></i>")
code("    </div>")
code("    <span>Manager Panel</span>")
code("  </div>")
code("  <nav class=\"sidebar-nav\">")
code("    <div class=\"nav-label\">Overview</div>")
code("    <a href=\"/manager/dashboard\" class=\"nav-link active\">")
code("      <i class=\"bi bi-grid\"></i> Dashboard")
code("    </a>")
code("    <a href=\"/manager/transactions\" class=\"nav-link\">")
code("      <i class=\"bi bi-arrow-left-right\"></i> All Transactions")
code("    </a>")
code("    <div class=\"nav-label\" style=\"margin-top:8px\">Analytics</div>")
code("    <a href=\"/manager/reports\" class=\"nav-link\">")
code("      <i class=\"bi bi-bar-chart-line\"></i> Reports")
code("    </a>")
code("  </nav>")
code("  <div class=\"sidebar-footer\">")
code("    <a href=\"/manager/logout\" class=\"nav-link\">")
code("      <i class=\"bi bi-box-arrow-left\"></i> Logout")
code("    </a>")
code("  </div>")
code("</aside>")

body("CSS-клас .mgr-sidebar задає фон #064e3b (темно-смарагдовий). "
     "Навігаційне меню містить лише 3 пункти: Dashboard, All Transactions, Reports. "
     "Відсутні пункти «Manage Users» та «Settings», що присутні у адміністратора. "
     "Клас .active встановлюється вручну на відповідному посиланні у кожному шаблоні.")

h3("Код топбара з manager/dashboard.ftl")
code("<div class=\"topbar\">")
code("  <h5>")
code("    <i class=\"bi bi-grid me-2\" style=\"color:#059669\"></i>")
code("    Manager Dashboard")
code("  </h5>")
code("  <div class=\"d-flex align-items-center gap-2\">")
code("    <div class=\"mgr-avatar\">")
code("      <#if manager??>")
code("        ${manager.firstName?substring(0,1)}${manager.lastName?substring(0,1)}")
code("      <#else>M</#if>")
code("    </div>")
code("    <div>")
code("      <div style=\"font-size:.85rem;font-weight:600;color:#111827\">")
code("        <#if manager??>${manager.firstName} ${manager.lastName}<#else>Manager</#if>")
code("      </div>")
code("      <div style=\"font-size:.7rem;color:#6b7280\">Manager</div>")
code("    </div>")
code("  </div>")
code("</div>")

body("CSS-клас .mgr-avatar використовує градієнт linear-gradient(135deg,#059669,#10b981) "
     "на відміну від .admin-avatar (#4f46e5,#7c3aed). Аватар відображає ініціали "
     "авторизованого менеджера через FreeMarker-директиву ?substring(0,1). "
     "Підпис ролі «Manager» є статичним текстом, що унеможливлює підміну в UI.")

h3("Спільна CSS-схема сайдбара")
code(".mgr-sidebar {")
code("  background: #064e3b;  /* темно-смарагд */")
code("}")
code(".mgr-avatar {")
code("  background: linear-gradient(135deg,#059669,#10b981);")
code("}")
code(".sidebar-nav .nav-link.active {")
code("  background: rgba(255,255,255,.18);  /* яскравіше ніж у admin (.15) */")
code("}")
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. ДЕТАЛЬНИЙ ОПИС КОЖНОГО ШАБЛОНУ
# ════════════════════════════════════════════════════════════
h1("3. Детальний опис кожного шаблону (manager/)")

# 3.1 login
h2("3.1. manager/login.ftl — Сторінка входу менеджера")

h3("Повна структура сторінки")
body("Шаблон login.ftl є самостійною сторінкою без сайдбара — застосовується "
     "окремий fullscreen layout. HTML-структура:")
for item in [
    "<!DOCTYPE html> / <html lang=\"en\"> / <head> — метатеги, Bootstrap 5.3.3 CDN, Inter шрифт, вбудований CSS.",
    "<body> — flex-контейнер (align-center, justify-center, min-height:100vh) з градієнтним фоном.",
    ".login-card — центральний білий блок (max-width:420px, border-radius:16px, padding:48px 40px).",
    "brand-icon block — іконка speedometer2 у зеленому квадраті.",
    "Заголовок «Manager Panel» та підзаголовок.",
    "FreeMarker-блок помилки (flash-повідомлення).",
    "HTML-форма з двома полями та кнопкою входу.",
    "Посилання «Back to PayFlow».",
    "<script> — JavaScript для toggle пароля.",
]:
    bullet(item)

h3("Всі змінні FreeMarker")
tbl(
    ["Змінна", "Тип", "Джерело", "Опис"],
    [
        ["error", "String", "RedirectAttributes flash", "Повідомлення про помилку входу. Передається при redirect після невдалої автентифікації. Значення: «User not found or insufficient privileges», «Incorrect password», «Your account has been blocked»"],
        ["RequestParameters.email", "String", "HTTP Request", "Значення параметра email з поточного або попереднього запиту. Використовується для pre-fill поля щоб не очищати email після помилки"],
    ],
    [3.5, 2.0, 3.0, 8.5]
)

h3("Відмінності від admin/login.ftl")
tbl(
    ["Параметр", "admin/login.ftl", "manager/login.ftl"],
    [
        ["CSS фон body",       "gradient #1e1b4b→#312e81→#1e1b4b", "gradient #022c22→#064e3b→#022c22"],
        ["brand-icon фон",     "gradient #4f46e5→#7c3aed",         "gradient #059669→#10b981"],
        ["Іконка",             "bi-shield-lock-fill",               "bi-speedometer2"],
        ["Заголовок h4",       "Admin Panel (color:#1e1b4b)",        "Manager Panel (color:#064e3b)"],
        ["form action",        "POST /admin/login",                  "POST /manager/login"],
        ["focus border-color", "#4f46e5",                            "#059669"],
        ["focus box-shadow",   "rgba(79,70,229,.15)",                "rgba(5,150,105,.15)"],
        ["btn-login фон",      "gradient #4f46e5→#7c3aed",          "gradient #059669→#10b981"],
        ["back-link a color",  "#4f46e5",                            "#059669"],
        ["email placeholder",  "admin@company.com",                  "manager@company.com"],
        ["Кнопка текст",       "Sign in to Admin Panel",             "Sign in to Manager Panel"],
        ["Додаткова перевірка","Немає (лише роль ADMIN)",            "Перевірка isActive у контролері"],
    ],
    [5.0, 5.5, 6.5]
)

h3("Повний код форми входу")
code("<form method=\"post\" action=\"/manager/login\" autocomplete=\"on\">")
code("  <div class=\"mb-3\">")
code("    <label class=\"form-label\" for=\"email\">Email address</label>")
code("    <div class=\"input-group\">")
code("      <span class=\"input-group-text\"><i class=\"bi bi-envelope\"></i></span>")
code("      <input type=\"email\" id=\"email\" name=\"email\"")
code("             class=\"form-control\"")
code("             placeholder=\"manager@company.com\" required autofocus")
code("             value=\"${(RequestParameters.email)!}\">")
code("    </div>")
code("  </div>")
code("  <div class=\"mb-4\">")
code("    <label class=\"form-label\" for=\"password\">Password</label>")
code("    <div class=\"input-group\">")
code("      <span class=\"input-group-text\"><i class=\"bi bi-lock\"></i></span>")
code("      <input type=\"password\" id=\"password\" name=\"password\"")
code("             class=\"form-control\"")
code("             placeholder=\"Enter your password\" required>")
code("      <button type=\"button\" id=\"togglePassword\"")
code("              class=\"btn btn-outline-secondary px-3\">")
code("        <i class=\"bi bi-eye\" id=\"eyeIcon\"></i>")
code("      </button>")
code("    </div>")
code("  </div>")
code("  <button type=\"submit\" class=\"btn btn-login mb-3\">")
code("    <i class=\"bi bi-box-arrow-in-right me-2\"></i>")
code("    Sign in to Manager Panel")
code("  </button>")
code("</form>")
body("Форма використовує метод POST на /manager/login. Поле email має атрибут "
     "autofocus для зручності. Значення поля email відновлюється після помилки "
     "через FreeMarker-вираз ${(RequestParameters.email)!}. "
     "Кнопка toggle видимості пароля керується JavaScript без бібліотек.")

# 3.2 dashboard
h2("3.2. manager/dashboard.ftl — Аналітична панель менеджера")

h3("Всі змінні Model (19 атрибутів)")
tbl(
    ["Змінна", "Тип Java", "Картка/Блок", "Опис"],
    [
        ["manager",              "User",          "Топбар, аватар",              "Авторизований менеджер — ім'я та ініціали"],
        ["totalUsers",           "Long",          "Stat-card #1",                "Загальна кількість зареєстрованих користувачів"],
        ["totalTransactions",    "Long",          "Stat-card #2",                "Загальна кількість транзакцій у системі"],
        ["pendingApprovals",     "Long",          "Stat-card #4",                "Кількість транзакцій зі статусом PENDING"],
        ["activeCards",          "Long",          "Stat-card #3",                "Кількість активних кредитних карток"],
        ["totalBalance",         "BigDecimal",    "Stat-card #5",                "Сума балансів усіх рахунків"],
        ["newUsersToday",        "Long",          "Stat-card #1 badge",          "Нові реєстрації сьогодні"],
        ["avgTransactionAmount", "BigDecimal",    "Stat-card #6",                "Середня сума завершених транзакцій"],
        ["completedCount",       "Long",          "Stat-card #7, графік",        "Кількість транзакцій COMPLETED"],
        ["pendingCount",         "Long",          "Графік (doughnut)",           "Кількість PENDING для легенди"],
        ["failedCount",          "Long",          "Stat-card #2 sub, графік",    "Кількість FAILED"],
        ["cancelledCount",       "Long",          "Графік (doughnut)",           "Кількість CANCELLED"],
        ["chartLabels",          "List<String>",  "volumeChart (лінійний)",      "Підписи осі X: останні 6 місяців (MMM yyyy)"],
        ["chartData",            "List<Long>",    "volumeChart (лінійний)",      "Кількість транзакцій за кожен місяць"],
        ["typePayment",          "Long",          "Type bars",                   "Кількість транзакцій типу PAYMENT"],
        ["typeReplenishment",    "Long",          "Type bars",                   "Кількість транзакцій типу REPLENISHMENT"],
        ["typeTransfer",         "Long",          "Type bars",                   "Кількість транзакцій типу TRANSFER"],
        ["topUsers",             "List<Map>",     "Top Users panel",             "Топ-5 користувачів: name, email, total, count, pct"],
        ["recentTransactions",   "List<Payment>", "Recent Transactions table",   "5 останніх транзакцій за датою (DESC)"],
    ],
    [4.0, 3.0, 4.0, 6.0]
)

h3("UI компоненти")
body("Дашборд менеджера містить наступні блоки:")
for item in [
    "readonly-banner (зелений) — відображається постійно на початку контентної зони.",
    "Рядок 1 (Row 1) — 4 первинні stat-card: Registered Users (+badge newUsersToday), "
    "Total Transactions (completed•failed), Active Cards, Pending Transactions.",
    "Рядок 2 (Row 2) — 4 вторинні stat-card: Total Balance ($), Avg Transaction ($), "
    "Completion Rate (%), System Health (99.8%).",
    "Рядок 3 (Row 3) — 2 панелі з графіками: volumeChart (line, col-lg-8) та "
    "statusChart з легендою та type-bars (col-lg-4).",
    "Рядок 4 (Row 4) — 2 панелі: Top Users (col-lg-5) та Recent Transactions (col-lg-7).",
]:
    bullet(item)

h3("Банер View-only mode — повний код")
code("<div class=\"readonly-banner\">")
code("  <i class=\"bi bi-eye\"></i>")
code("  <span style=\"font-size:.85rem;color:#065f46\">")
code("    <strong>View-only mode.</strong>")
code("    You can monitor system statistics and transactions.")
code("    Contact an administrator for any changes.")
code("  </span>")
code("</div>")
body("CSS-клас .readonly-banner реалізований наступним чином:")
code(".readonly-banner {")
code("  background: linear-gradient(135deg,#ecfdf5,#d1fae5);")
code("  border: 1px solid #6ee7b7;")
code("  border-radius: 12px;")
code("  padding: 12px 18px;")
code("  margin-bottom: 24px;")
code("  display: flex;")
code("  align-items: center;")
code("  gap: 12px;")
code("}")
body("Банер відображається без жодних умов — він присутній на сторінці завжди, "
     "на відміну від адмін-панелі де pending-banner показується лише при "
     "pendingApprovals > 0. Це принципова різниця: у менеджера статус «тільки "
     "перегляд» постійний, а у адміністратора — попередження ситуативне.")

h3("Всі FreeMarker директиви з прикладами коду")
body("Ключові FreeMarker-конструкції у dashboard.ftl:")

code("<!-- Аватар з ініціалами та fallback -->")
code("<#if manager??>")
code("  ${manager.firstName?substring(0,1)}${manager.lastName?substring(0,1)}")
code("<#else>M</#if>")

code("<!-- Умовний badge «+N today» -->")
code("<#if (newUsersToday > 0)>")
code("  <span class=\"bdg bdg-teal\">+${newUsersToday} today</span>")
code("</#if>")

code("<!-- Обчислення відсотка завершення -->")
code("<#assign successRate = 0>")
code("<#if (totalTransactions > 0)>")
code("  <#assign successRate = (completedCount * 100 / totalTransactions)>")
code("</#if>")
code("${successRate?string[\"0\"]}%")

code("<!-- Горизонтальний бар типів транзакцій -->")
code("<#assign totalTx = typePayment + typeReplenishment + typeTransfer>")
code("<#if (totalTx == 0)><#assign totalTx = 1></#if>")
code("width:${(typePayment*100/totalTx)?string[\"0\"]}%")
code("width:${(typeReplenishment*100/totalTx)?string[\"0\"]}%")
code("width:${(typeTransfer*100/totalTx)?string[\"0\"]}%")

code("<!-- Топ-5 користувачів з рангом та роздільником -->")
code("<#list topUsers as u>")
code("  <#assign rank = u_index + 1>")
code("  <div style=\"padding:14px 20px;")
code("    border-bottom:<#if u_has_next>1px solid #f9fafb<#else>none</#if>\">")
code("    <div style=\"background:")
code("      <#if rank==1>#fef3c7<#elseif rank==2>#f3f4f6<#else>#f9fafb</#if>;")
code("      color:")
code("      <#if rank==1>#b45309<#elseif rank==2>#374151<#else>#9ca3af</#if>\">")
code("      ${rank}")
code("    </div>")
code("    ${u.name} &bull; ${u.count} txn<#if (u.count > 1)>s</#if>")
code("    $${u.total?string[\"0.00\"]}")
code("    width:${u.pct?string[\"0.0\"]}%")
code("  </div>")
code("</#list>")

code("<!-- Таблиця останніх транзакцій з кольором суми -->")
code("<#list recentTransactions as p>")
code("  <td><code style=\"color:#059669\">${p.transactionId!\"—\"}</code></td>")
code("  <td>")
code("    <#if p.createdAt??>")
code("      <#assign mo=[\"Jan\",\"Feb\",\"Mar\",\"Apr\",\"May\",\"Jun\",")
code("                    \"Jul\",\"Aug\",\"Sep\",\"Oct\",\"Nov\",\"Dec\"]>")
code("      ${mo[p.createdAt.monthValue-1]} ${p.createdAt.dayOfMonth}")
code("    <#else>—</#if>")
code("  </td>")
code("  <td>")
code("    <#if p.type?? && p.type.name()==\"REPLENISHMENT\">")
code("      <span class=\"amt-pos\">+$${p.amount?string[\"0.00\"]}</span>")
code("    <#elseif p.type?? && p.type.name()==\"TRANSFER\">")
code("      <span class=\"amt-neu\">$${p.amount?string[\"0.00\"]}</span>")
code("    <#else>")
code("      <span class=\"amt-neg\">-$${p.amount?string[\"0.00\"]}</span>")
code("    </#if>")
code("  </td>")
code("</#list>")

code("<!-- Chart.js дані генеруються через FreeMarker -->")
code("const volumeLabels =")
code("  [<#list chartLabels as l>\"${l}\"<#sep>,</#list>];")
code("const volumeData =")
code("  [<#list chartData as d>${d}<#sep>,</#list>];")
code("const statusData =")
code("  [${completedCount},${pendingCount},${failedCount},${cancelledCount}];")

# 3.3 transactions
h2("3.3. manager/transactions.ftl — Перегляд транзакцій")

h3("Структура таблиці транзакцій")
body("Таблиця транзакцій менеджера містить 7 колонок (на відміну від 8 в адміна):")
tbl(
    ["#", "Колонка", "Дані", "Особливості"],
    [
        ["1", "TXN ID",      "p.transactionId",                    "Monospace, колір #059669 (смарагд замість індиго #4f46e5)"],
        ["2", "Date / Time", "p.createdAt (дата + час)",           "Масив mo[] для назви місяця, формат HH:MM"],
        ["3", "User",        "p.account.creditCard.user (ланцюжок)", "Ім'я жирним + email сірим. Null-safe: 3 рівні перевірки ??"],
        ["4", "Description", "p.description",                      "Truncation: max-width:160px, text-overflow:ellipsis"],
        ["5", "Type",        "p.type.name()",                      "Кольорові бейджі: bdg-replenishment/transfer/payment"],
        ["6", "Amount",      "p.amount",                           "Кольорове форматування: amt-pos(+)/amt-neg(-)/amt-neu($)"],
        ["7", "Status",      "p.status.name()",                    "Кольорові бейджі: completed/pending/failed/cancelled"],
    ],
    [1.0, 3.0, 4.5, 8.5]
)
body("Критична відмінність від admin/transactions.ftl: колонка «Actions» (8-ма) "
     "повністю відсутня. Немає ні заголовка, ні клітинок, ні форм POST. "
     "Empty-state colspan дорівнює 7 (замість 8 в адміна).")

h3("Фільтри та пагінація — FTL код")
body("Фільтрація виконується на сервері через GET-форму:")
code("<form method=\"get\" action=\"/manager/transactions\"")
code("      class=\"table-card mb-0\"")
code("      style=\"border-radius:14px 14px 0 0;border-bottom:none\">")
code("  <div class=\"table-card-header\">")
code("    <!-- Пошуковий рядок -->")
code("    <div class=\"search-wrapper\">")
code("      <i class=\"bi bi-search\"></i>")
code("      <input type=\"text\" name=\"search\"")
code("             value=\"${searchQuery!\"\"}\"")
code("             placeholder=\"Search by description or TXN ID…\">")
code("    </div>")
code("    <!-- Select типу транзакції -->")
code("    <select name=\"type\">")
code("      <option value=\"ALL\"")
code("        <#if selectedType==\"ALL\">selected</#if>>All Types</option>")
code("      <option value=\"PAYMENT\"")
code("        <#if selectedType==\"PAYMENT\">selected</#if>>Payment</option>")
code("      <option value=\"REPLENISHMENT\"")
code("        <#if selectedType==\"REPLENISHMENT\">selected</#if>>Top-up</option>")
code("      <option value=\"TRANSFER\"")
code("        <#if selectedType==\"TRANSFER\">selected</#if>>Transfer</option>")
code("    </select>")
code("    <!-- Select статусу -->")
code("    <select name=\"status\">")
code("      <option value=\"ALL\"")
code("        <#if selectedStatus==\"ALL\">selected</#if>>All Statuses</option>")
code("      <option value=\"COMPLETED\"")
code("        <#if selectedStatus==\"COMPLETED\">selected</#if>>Completed</option>")
code("      <option value=\"PENDING\"")
code("        <#if selectedStatus==\"PENDING\">selected</#if>>Pending</option>")
code("      <option value=\"FAILED\"")
code("        <#if selectedStatus==\"FAILED\">selected</#if>>Failed</option>")
code("      <option value=\"CANCELLED\"")
code("        <#if selectedStatus==\"CANCELLED\">selected</#if>>Cancelled</option>")
code("    </select>")
code("    <input type=\"hidden\" name=\"page\" value=\"0\">")
code("    <!-- Кнопка Filter (зелена) -->")
code("    <button type=\"submit\"")
code("      style=\"background:#059669;color:#fff;\">")
code("      <i class=\"bi bi-funnel me-1\"></i>Filter")
code("    </button>")
code("    <!-- Кнопка Reset -->")
code("    <a href=\"/manager/transactions\"")
code("       class=\"btn btn-sm btn-light\">Reset</a>")
code("  </div>")
code("</form>")

body("Пагінація — smart-алгоритм з FreeMarker:")
code("<#if (totalPages > 1)>")
code("  <span>Page ${currentPage + 1} of ${totalPages}")
code("    &bull; ${totalItems} total</span>")
code("  <!-- Кнопка «попередня» -->")
code("  <a href=\"/manager/transactions?page=${currentPage-1}")
code("       &search=${searchQuery!\"\"}&type=${selectedType!\"\"}")
code("       &status=${selectedStatus!\"\"}\"")
code("     class=\"page-btn <#if !hasPrevious>disabled</#if>\">")
code("    <i class=\"bi bi-chevron-left\"></i>")
code("  </a>")
code("  <!-- Номери сторінок з '...' -->")
code("  <#list 0..(totalPages - 1) as i>")
code("    <#if (i == 0) || (i == totalPages - 1)")
code("         || ((i >= currentPage - 1)")
code("         && (i <= currentPage + 1))>")
code("      <a href=\"...?page=${i}...\"")
code("         class=\"page-btn <#if i == currentPage>active</#if>\">")
code("        ${i + 1}")
code("      </a>")
code("    <#elseif (i == 1 && currentPage > 3)")
code("         || (i == totalPages-2 && currentPage < totalPages-4)>")
code("      <span class=\"page-btn\">…</span>")
code("    </#if>")
code("  </#list>")
code("  <!-- Кнопка «наступна» -->")
code("  <a href=\"/manager/transactions?page=${currentPage+1}...\"")
code("     class=\"page-btn <#if !hasNext>disabled</#if>\">")
code("    <i class=\"bi bi-chevron-right\"></i>")
code("  </a>")
code("</#if>")
body("Алгоритм пагінації завжди відображає першу та останню сторінку, "
     "а також ±1 від поточної. Між ними вставляє «…» при необхідності. "
     "Параметри фільтрів зберігаються у URL при переході між сторінками.")

h3("Відсутні елементи порівняно з admin/transactions.ftl")
tbl(
    ["Елемент", "admin/transactions.ftl", "manager/transactions.ftl"],
    [
        ["Колонка «Actions»",         "Присутня (8-ма колонка)",             "Відсутня"],
        ["Форма approve",             "POST /admin/transactions/approve",     "Відсутня"],
        ["Форма reject",              "POST /admin/transactions/reject",      "Відсутня"],
        ["<#assign isPending ...>",   "Присутнє для умови відображення кнопок","Відсутнє"],
        ["btn-approve / btn-reject",  "CSS-класи з кольоровими стилями",     "Відсутні"],
        ["Колір TXN ID code",         "color:#4f46e5 (індиго)",              "color:#059669 (смарагд)"],
        ["Колір кнопки Filter",       "background:#4f46e5",                  "background:#059669"],
        ["Активна сторінка пагінації","background:#4f46e5",                  "background:#059669"],
        ["focus-color input/select",  "border-color:#4f46e5",                "border-color:#059669"],
        ["colspan empty-state",       "colspan=\"8\"",                       "colspan=\"7\""],
        ["Readonly-notice badge",     "Відсутній",                           "Присутній (зелений)"],
        ["Flash-alerts",              "successMessage / errorMessage",        "Відсутні (немає дій)"],
    ],
    [5.5, 5.5, 6.0]
)

h3("Код readonly-notice")
code("<div class=\"d-flex align-items-center justify-content-between mb-3\">")
code("  <div>")
code("    <h4 class=\"section-title\">All Transactions</h4>")
code("    <p class=\"text-muted small mb-0 mt-1\">")
code("      ${totalItems} total records")
code("    </p>")
code("  </div>")
code("  <!-- readonly-notice badge -->")
code("  <div class=\"readonly-notice\">")
code("    <i class=\"bi bi-eye\"></i>")
code("    View-only — no approval actions available")
code("  </div>")
code("</div>")
body("CSS-клас .readonly-notice:")
code(".readonly-notice {")
code("  background: #ecfdf5;")
code("  border: 1px solid #6ee7b7;")
code("  border-radius: 10px;")
code("  padding: 10px 16px;")
code("  font-size: .82rem;")
code("  color: #065f46;")
code("  display: flex;")
code("  align-items: center;")
code("  gap: 8px;")
code("}")
body("Readonly-notice відображається у правій частині заголовного рядка сторінки. "
     "На відміну від readonly-banner на дашборді (більший, повноширинний), "
     "цей елемент компактний і не займає окремий рядок.")

# 3.4 reports
h2("3.4. manager/reports.ftl — Заглушка сторінки звітів")

h3("Структура placeholder сторінки")
body("Шаблон reports.ftl реалізує повноцінний layout панелі (сайдбар + топбар) "
     "з єдиним вертикально відцентрованим блоком-заглушкою. Структура контентної зони:")
for item in [
    "Div .content-area з класами d-flex align-items-center justify-content-center та min-height:calc(100vh - 65px) — вертикальне центрування на всю висоту екрана мінус висота топбара.",
    "Внутрішній div max-width:400px з класом text-center.",
    "Іконка-блок 80×80px: background:#d1fae5, border-radius:20px, іконка bi-bar-chart-line (2rem, color:#059669).",
    "Заголовок h4 «Reports Coming Soon» (font-weight:700, color:#111827).",
    "Параграф з поясненням що функціонал у розробці та рекомендацією перейти на Dashboard.",
    "Кнопка «Go to Dashboard» — посилання на /manager/dashboard з зеленим фоном (#059669).",
]:
    bullet(item)

h3("FreeMarker-змінні та директиви")
body("Шаблон використовує лише одну змінну manager (User) для відображення "
     "авторизованого менеджера у топбарі. Директиви шаблону мінімальні:")
code("<#-- @ftlvariable name=\"manager\"")
code("     type=\"ua.com.kisit.course2026np.entity.User\" -->")
code("")
code("<!-- Ініціали в аватарі топбара -->")
code("<#if manager?>")
code("  ${manager.firstName?substring(0,1)}")
code("  ${manager.lastName?substring(0,1)}")
code("<#else>M</#if>")
code("")
code("<!-- Повне ім'я у топбарі -->")
code("<#if manager?>")
code("  ${manager.firstName} ${manager.lastName}")
code("<#else>Manager</#if>")

h3("Що планується реалізувати у майбутньому")
body("Placeholder-текст шаблону містить опис: «Advanced analytics and export "
     "features are in development. Check the Dashboard for real-time metrics.» "
     "Виходячи з загальної архітектури застосунку та наявних даних у моделі, "
     "планований функціонал може включати:")
for item in [
    "Звіт по транзакціях за довільний період із вибором дат (date range picker).",
    "Аналіз за типами транзакцій (Payment/Replenishment/Transfer) у вигляді таблиці та графіків.",
    "Звіт по активності користувачів: реєстрації, транзакції, баланси.",
    "Експорт даних у форматах CSV та PDF для зовнішнього аналізу.",
    "Порівняльна аналітика: поточний місяць vs попередній, динаміка показників.",
]:
    bullet(item)
body("Відмінності від admin/reports.ftl: зелений фон іконки (#d1fae5 замість #ede9fe), "
     "зелений колір іконки (#059669 замість #7c3aed), зелена кнопка (#059669 замість #4f46e5), "
     "посилання на /manager/dashboard замість /admin/dashboard.")
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. МЕХАНІЗМ ЗАХИСТУ ПРЕДСТАВЛЕНЬ
# ════════════════════════════════════════════════════════════
h1("4. Механізм захисту представлень")

h2("4.1. Метод isNotAuthenticated() у ManagerController")
body("Захист доступу до шаблонів менеджерської панелі реалізований через "
     "HttpSession-перевірку у кожному методі ManagerController. Допоміжний "
     "приватний метод isNotAuthenticated() перевіряє наявність сесійного атрибута:")
code("private boolean isNotAuthenticated(HttpSession session) {")
code("    return session.getAttribute(\"managerUser\") == null;")
code("}")
body("Цей метод викликається на початку КОЖНОГО обробника контролера:")
code("@GetMapping(\"/dashboard\")")
code("public String dashboard(Model model, HttpSession session) {")
code("    if (isNotAuthenticated(session))")
code("        return \"redirect:/manager/login\";")
code("    // ... логіка ...")
code("    return \"manager/dashboard\";")
code("}")
body("Таким чином, якщо сесійний атрибут «managerUser» відсутній — контролер "
     "повертає redirect, а не назву шаблону. FreeMarker-рендеринг не запускається. "
     "Шаблон фізично недоступний без авторизації.")

h2("4.2. Що відбувається при несанкціонованому доступі")
body("Розглянемо сценарії спроби доступу до /manager/* для різних ролей:")
tbl(
    ["Роль / Стан", "Атрибут сесії", "Спроба GET /manager/dashboard", "Результат"],
    [
        ["Не авторизований", "Немає", "isNotAuthenticated() = true", "redirect:/manager/login"],
        ["CLIENT (авторизований)", "userId + userRole=CLIENT", "isNotAuthenticated() = true (немає managerUser)", "redirect:/manager/login"],
        ["MANAGER (авторизований)", "managerUser = User", "isNotAuthenticated() = false", "Рендеринг manager/dashboard.ftl"],
        ["MANAGER (заблокований)", "Немає (відмова при login)", "Не може авторизуватись", "redirect:/manager/login з «Your account has been blocked»"],
        ["ADMIN через /admin/login", "adminUser = User", "isNotAuthenticated() = true (немає managerUser)", "redirect:/manager/login"],
        ["ADMIN через /login", "adminUser + userId", "isNotAuthenticated() = true (немає managerUser)", "redirect:/manager/login"],
    ],
    [4.0, 4.0, 5.5, 3.5]
)
body("Ключова особливість: навіть авторизований ADMIN не може отримати доступ "
     "до менеджерських сторінок без атрибута «managerUser» у своїй сесії. "
     "Сесії admin і manager повністю розділені і не перетинаються.")

h2("4.3. Перевірка isActive при вході менеджера")
body("На відміну від AdminController, ManagerController містить додаткову перевірку "
     "активності акаунта при вході:")
code("@PostMapping(\"/login\")")
code("public String handleLogin(")
code("        @RequestParam String email,")
code("        @RequestParam String password,")
code("        HttpSession session,")
code("        RedirectAttributes redirectAttributes) {")
code("")
code("    Optional<User> managerOpt = userRepository.findAll().stream()")
code("        .filter(u -> u.getEmail()")
code("                      .equalsIgnoreCase(email.trim())")
code("                   && u.getRole() == UserRole.MANAGER)")
code("        .findFirst();")
code("")
code("    if (managerOpt.isEmpty()) {")
code("        redirectAttributes.addFlashAttribute(\"error\",")
code("            \"User not found or insufficient privileges\");")
code("        return \"redirect:/manager/login\";")
code("    }")
code("")
code("    User manager = managerOpt.get();")
code("    if (!manager.getPassword().equals(password.trim())) {")
code("        redirectAttributes.addFlashAttribute(\"error\",")
code("            \"Incorrect password\");")
code("        return \"redirect:/manager/login\";")
code("    }")
code("")
code("    // Додаткова перевірка — відсутня в AdminController")
code("    if (!Boolean.TRUE.equals(manager.getIsActive())) {")
code("        redirectAttributes.addFlashAttribute(\"error\",")
code("            \"Your account has been blocked\");")
code("        return \"redirect:/manager/login\";")
code("    }")
code("")
code("    session.setAttribute(\"managerUser\", manager);")
code("    return \"redirect:/manager/dashboard\";")
code("}")
body("Послідовність перевірок: 1) пошук користувача з роллю MANAGER за email; "
     "2) перевірка пароля; 3) перевірка isActive. Лише після успішного проходження "
     "всіх трьох перевірок об'єкт User зберігається в сесії. Адміністратор може "
     "заблокувати менеджера через /admin/users, встановивши isActive=false — "
     "і менеджер більше не зможе увійти без повторного розблокування.")
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. АНАЛІЗ ВІДПОВІДНОСТІ ВИМОГАМ
# ════════════════════════════════════════════════════════════
h1("5. Аналіз відповідності вимогам")
tbl(
    ["Вимога", "Реалізація у шаблонах", "Статус"],
    [
        ["Окрема сторінка входу для менеджера",
         "manager/login.ftl з унікальним дизайном (зелена гама, speedometer icon)",
         "Виконано"],
        ["Аналітичний дашборд для менеджера",
         "manager/dashboard.ftl — 8 stat-card, 2 Chart.js графіки, топ-5, recent transactions",
         "Виконано"],
        ["Режим «тільки перегляд» — явна комунікація",
         "readonly-banner (dashboard) та readonly-notice (transactions) у зеленому стилі",
         "Виконано"],
        ["Перегляд транзакцій з фільтрацією",
         "manager/transactions.ftl — GET-форма: search + type-select + status-select",
         "Виконано"],
        ["Пагінація у таблиці транзакцій",
         "Smart-алгоритм <#list 0..(totalPages-1) as i> з '…' та збереженням query params",
         "Виконано"],
        ["Відсутність кнопок approve/reject у менеджера",
         "Колонка Actions фізично відсутня у manager/transactions.ftl (7 колонок)",
         "Виконано"],
        ["Сторінка звітів",
         "manager/reports.ftl — placeholder з описом майбутнього функціоналу",
         "Частково (заглушка)"],
        ["Колірна диференціація від admin панелі",
         "Смарагдовий #064e3b/#059669 vs індиго #1e1b4b/#4f46e5 — всі 4 шаблони",
         "Виконано"],
        ["Захист від несанкціонованого доступу",
         "isNotAuthenticated() на початку кожного методу ManagerController",
         "Виконано"],
        ["Перевірка isActive при вході",
         "handleLogin() перевіряє !manager.getIsActive() перед збереженням в сесію",
         "Виконано"],
        ["Відображення імені менеджера у топбарі",
         "${manager.firstName} ${manager.lastName} у всіх шаблонах",
         "Виконано"],
        ["Null-safe виведення даних",
         "variable?? та !\"default\" у всіх 4 шаблонах, ланцюжок ?? у user-навігації",
         "Виконано"],
        ["Налаштування профілю менеджера",
         "Сторінка settings для менеджера відсутня",
         "Не реалізовано"],
        ["Адаптивна верстка",
         ".col-xl-3.col-md-6 для stat-card, .col-lg-8/4 та .col-lg-5/7 для панелей",
         "Виконано"],
        ["Підтримка порожніх станів",
         "Empty-state блоки з bi-inbox іконкою при відсутності даних у таблицях",
         "Виконано"],
    ],
    [5.5, 8.0, 3.5]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. ВИСНОВКИ
# ════════════════════════════════════════════════════════════
h1("6. Висновки")
body("У ході виконання лабораторної роботи №14 було детально проаналізовано та "
     "описано чотири FreeMarker-шаблони менеджерської панелі платіжного застосунку "
     "PayFlow (manager/login.ftl, manager/dashboard.ftl, manager/transactions.ftl, "
     "manager/reports.ftl) та контролер ManagerController.java.")
body("Менеджерська панель реалізує принцип найменших привілеїв: менеджер отримує "
     "доступ лише до трьох розділів (Dashboard, All Transactions, Reports), "
     "кожен з яких доступний виключно у режимі перегляду. Відсутність кнопок "
     "дій та форм не є наслідком CSS-приховування — відповідні HTML-елементи "
     "просто відсутні у шаблонах. Це надійніший підхід з точки зору безпеки.")
body("Ключові технічні рішення: readonly-banner на дашборді та readonly-notice на "
     "сторінці транзакцій чітко комунікують обмежений статус ролі. Зелена колірна "
     "схема (#064e3b для сайдбара, #059669 для акцентів) візуально відрізняє "
     "менеджерську панель від адміністраторської (темно-індиго #1e1b4b / #4f46e5) "
     "без текстових підказок — достатньо поглянути на колір сайдбара.")
body("Механізм захисту побудований на HttpSession-перевірці: метод "
     "isNotAuthenticated() перевіряє наявність атрибута «managerUser» на початку "
     "кожного обробника. Додатково при вході перевіряється isActive — "
     "адміністратор може оперативно відключити будь-якого менеджера, і той "
     "більше не зможе авторизуватись. Сесії admin і manager повністю ізольовані.")
body("FreeMarker-шаблони активно використовують 18 різних директив та конструкцій: "
     "умовні блоки <#if>, цикли <#list> з item_index/item_has_next, локальні "
     "змінні <#assign>, форматування чисел ?string[\"0.00\"], null-safe оператори "
     "variable?? та !\"default\", а також генерацію JavaScript-масивів для "
     "Chart.js безпосередньо у шаблоні.")
body("Виявлено два технічних борги: відсутність сторінки налаштувань для менеджера "
     "та відсутність механізму include/extends (сайдбар та топбар дублюються "
     "у кожному з 4 шаблонів). Сторінка звітів є заглушкою, що відповідає "
     "поточному стану розробки застосунку.")

out = r"C:\Users\Lenovo\OneDrive\Desktop\practica\ЛР_14_Зівак_Сергій_371.docx"
doc.save(out)
print("Saved:", out)
