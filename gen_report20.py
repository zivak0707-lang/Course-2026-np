"""
Генератор лабораторної роботи №20 — Зівак Сергій, група 371
Тема: Запуск, тестування та профілювання застосунку PayFlow
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── PAGE MARGINS ────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── DEFAULT PARAGRAPH STYLE ─────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(0)


def set_font(run, size=14, bold=False, italic=False, color=None):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=14, bold=False, italic=False,
                  sb=0, sa=0, first_line=False, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(1.25)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(text, level=1, sb=12, sa=6):
    size = 16 if level == 1 else 14
    p = add_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=size, bold=True, sb=sb, sa=sa)
    return p


def add_subheading(text, sb=10, sa=4):
    p = add_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=14, bold=True, sb=sb, sa=sa)
    return p


def pc(text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4):
    return add_paragraph(text, align=align, size=size, bold=bold, sb=sb, sa=sa)


fig_counter = [0]


def add_screenshot_placeholder(caption: str):
    fig_counter[0] += 1
    n = fig_counter[0]

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    cell.width = Cm(15.5)

    # grey background
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D9D9D9')
    tc_pr.append(shd)

    # height of cell ≈ 8 cm via trHeight
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'),  str(int(Cm(8).pt * 20)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

    # centered placeholder text
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(0)
    cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    run = cp.add_run(f"[ МІСЦЕ ДЛЯ СКРІНШОТУ ]")
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # caption paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(10)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = cap.add_run(f"Рисунок {n} — {caption}")
    r.font.name   = 'Times New Roman'
    r.font.size   = Pt(12)
    r.font.italic = True


def body(text, first_line=True):
    add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=14, first_line=first_line)


# ════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════
pc("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ", 14, bold=True, sb=0, sa=2)
pc("КИЇВСЬКИЙ НАЦІОНАЛЬНИЙ ЕКОНОМІЧНИЙ УНІВЕРСИТЕТ", 14, bold=True, sa=2)
pc("імені Вадима Гетьмана", 14, sa=2)
pc("Факультет комп'ютерних наук та інформаційних технологій (ФКІСІТ)", 14, sa=60)

pc("ЗВІТ", 16, bold=True, sa=4)
pc("з лабораторної роботи №20", 14, bold=True, sa=4)
pc("Тема: «Запуск, тестування та профілювання застосунку»", 14, bold=True, sa=60)

pc("Виконав: студент групи 371", 14, sa=2)
pc("Зівак Сергій", 14, bold=True, sa=60)

pc("Київ — 2025", 14, sa=0)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  1. МЕТА РОБОТИ
# ════════════════════════════════════════════════════════════
add_heading("1. Мета роботи")
body("Метою лабораторної роботи є практичне освоєння навичок запуску, "
     "тестування та базового профілювання веб-застосунку PayFlow — "
     "навчальної банківської платіжної системи, реалізованої на платформі "
     "Spring Boot. У ході роботи студент запускає застосунок у локальному "
     "середовищі, перевіряє функціонал усіх ролей системи (CLIENT, MANAGER, "
     "ADMIN), аналізує лог-файли, а також досліджує особливості реалізації "
     "двох ключових фінансових операцій — скасування та повернення транзакцій.")


# ════════════════════════════════════════════════════════════
#  2. ТЕОРЕТИЧНА ЧАСТИНА
# ════════════════════════════════════════════════════════════
add_heading("2. Теоретична частина")

add_subheading("2.1 Життєвий цикл Spring Boot застосунку")
body("Spring Boot надає механізм автоматичної конфігурації (auto-configuration), "
     "який суттєво спрощує розгортання Java-застосунків. При старті застосунок "
     "проходить кілька фаз: завантаження середовища (Environment), "
     "ініціалізація контексту Spring (ApplicationContext), реєстрація бінів, "
     "виконання CommandLineRunner / ApplicationRunner, після чого сервер стає "
     "готовим до прийому HTTP-запитів. Вбудований контейнер Tomcat дозволяє "
     "запускати застосунок без зовнішнього сервлет-контейнера — достатньо "
     "виконати java -jar або запустити основний клас через IDE.")
body("Spring Boot 4.x (на базі Spring Framework 6) вимагає Java 17+ та "
     "використовує Jakarta EE 10 (простір імен jakarta.*). "
     "Конфігурація застосунку зберігається у файлі application.properties "
     "або application.yml, де задаються параметри підключення до бази даних, "
     "порт сервера, налаштування JPA та шаблонізатора FreeMarker.")

add_subheading("2.2 Підходи до тестування веб-застосунків")
body("Сучасна практика розробки передбачає кілька рівнів тестування. "
     "Модульні тести (unit tests) перевіряють окремі класи або методи в "
     "ізоляції від зовнішніх залежностей. Інтеграційні тести охоплюють "
     "взаємодію між компонентами системи — контролерами, сервісами та "
     "репозиторіями — і, як правило, потребують реальної або in-memory бази "
     "даних. У Spring Boot для інтеграційних тестів використовуються анотації "
     "@SpringBootTest та @DataJpaTest. Функціональне (ручне) тестування "
     "передбачає перевірку сценаріїв взаємодії користувача з інтерфейсом "
     "через браузер і є невід'ємною частиною перевірки коректності UI.")
body("Окремою категорією є тестування безпеки: перевірка захисту маршрутів "
     "від неавторизованого доступу, коректність перенаправлень при відсутній "
     "сесії (HTTP 401), а також перевірка захисту від несанкціонованих дій "
     "з чужими ресурсами (HTTP 403). Для банківських систем критично важливо "
     "тестувати граничні випадки фінансових операцій: від'ємний баланс, "
     "блокований рахунок, спроба повторного повернення вже поверненої транзакції.")

add_subheading("2.3 Профілювання JVM-застосунків")
body("Профілювання дозволяє виявити вузькі місця у продуктивності Java-застосунків. "
     "Інструменти JVisualVM, JMC (Java Mission Control), Async Profiler та "
     "Actuator від Spring Boot надають інформацію про завантаження CPU, "
     "споживання пам'яті heap/non-heap, активність garbage collector, а також "
     "кількість відкритих з'єднань з базою даних. У рамках навчального "
     "застосунку базовим інструментом профілювання слугує аналіз "
     "лог-файлів payflow.log та security.log, що дозволяє відстежити "
     "послідовність операцій та виявити потенційні помилки.")

add_subheading("2.4 Управління транзакціями у банківських системах")
body("У реляційних СУБД транзакція — це атомарна одиниця роботи, яка "
     "виконується повністю або не виконується взагалі (принцип ACID). "
     "Spring Framework забезпечує декларативне управління транзакціями через "
     "анотацію @Transactional, яка автоматично відкриває транзакцію на початку "
     "методу та фіксує або відкочує її залежно від результату виконання. "
     "Для банківських систем особливо важливим є аудит-трейл — незмінний "
     "журнал усіх фінансових операцій, що забезпечує відстежуваність і "
     "відповідність вимогам регуляторів. У системі PayFlow оригінальна "
     "транзакція ніколи не змінюється після завершення — замість цього "
     "створюється нова транзакція повернення, що зберігає повну історію "
     "всіх фінансових рухів.")


# ════════════════════════════════════════════════════════════
#  3. ОПИС ЗАСТОСУНКУ PAYFLOW
# ════════════════════════════════════════════════════════════
add_heading("3. Опис застосунку PayFlow")

add_subheading("3.1 Технічний стек")
body("PayFlow — навчальна банківська платіжна система, побудована на "
     "наступних технологіях:")

tech_items = [
    ("Spring Boot 4.0.2",
     "фреймворк для створення веб-застосунків на Java 17, забезпечує "
     "автоматичну конфігурацію та вбудований сервер Tomcat"),
    ("Spring Security",
     "фреймворк безпеки з підтримкою BCrypt-хешування паролів та PIN-кодів "
     "(strength=10), захист від brute-force атак через LoginAttemptService"),
    ("Spring Data JPA / Hibernate",
     "ORM-рівень для роботи з реляційною базою даних через репозиторії; "
     "ddl-auto=update для автоматичного оновлення схеми"),
    ("MySQL",
     "реляційна СУБД, підключення через spring.datasource.url; "
     "база даних course-2026-np"),
    ("FreeMarker",
     "серверний шаблонізатор для рендерингу HTML-сторінок (.ftl файли); "
     "шаблони розділені по папках admin/, manager/, client/"),
    ("Bootstrap 5",
     "CSS-фреймворк для адаптивного інтерфейсу"),
    ("Chart.js",
     "JavaScript-бібліотека для побудови графіків у дашборді менеджера"),
    ("BCrypt (Spring Security Crypto)",
     "алгоритм хешування для паролів та PIN-кодів; "
     "окремий BCryptPasswordEncoder(10) для PIN"),
    ("SLF4J + Logback",
     "система логування з двома окремими файлами: "
     "payflow.log (загальний) та security.log (аудит безпеки)"),
    ("Lombok",
     "генерація boilerplate-коду: @Builder, @Getter, @Setter, @Slf4j"),
    ("Jakarta Validation",
     "валідація вхідних даних через @NotBlank, @Pattern та ін."),
]

for tech, desc in tech_items:
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(f"• {tech}: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.bold = True
    r2 = p.add_run(desc)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)

add_subheading("3.2 Ролі системи та їх можливості")

body("Система PayFlow реалізує три ролі користувачів з різними рівнями доступу:")

# ─ CLIENT ─
p = add_paragraph("Роль CLIENT (клієнт):", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=8, sa=2)
client_funcs = [
    "Перегляд власного дашборду з балансом та останніми транзакціями",
    "Управління рахунками: перегляд, створення нових рахунків",
    "Управління картками: перегляд карток, блокування/розблокування",
    "Встановлення денного ліміту витрат на картці (spendingLimit)",
    "Зміна PIN-коду картки (4 цифри, BCrypt-хешування)",
    "Переказ коштів між рахунками (PaymentType.TRANSFER)",
    "Здійснення платежів (PaymentType.PAYMENT)",
    "Поповнення рахунку (PaymentType.REPLENISHMENT)",
    "Перегляд власної історії транзакцій",
    "Зміна пароля особистого кабінету",
]
for f in client_funcs:
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent  = Cm(1.25)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p2.add_run(f"— {f}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

# ─ MANAGER ─
p = add_paragraph("Роль MANAGER (менеджер):", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=8, sa=2)
manager_funcs = [
    "Перегляд дашборду зі статистикою по всіх транзакціях системи",
    "Моніторинг усіх транзакцій системи (тільки читання)",
    "Перегляд аналітики з графіками Chart.js (суми по типах, динаміка)",
    "Перегляд статистики (загальна сума, кількість транзакцій, середнє)",
]
body("Менеджер має лише права на читання — жодних дій зі зміни даних не виконує.")
for f in manager_funcs:
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent  = Cm(1.25)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p2.add_run(f"— {f}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

# ─ ADMIN ─
p = add_paragraph("Роль ADMIN (адміністратор):", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=8, sa=2)
admin_funcs = [
    "Перегляд усіх користувачів системи, їх блокування/розблокування",
    "Призначення/зміна ролей користувачів",
    "Перегляд та управління всіма транзакціями системи",
    "Підтвердження транзакцій (зміна статусу PENDING → COMPLETED)",
    "Відхилення транзакцій (зміна статусу PENDING → FAILED)",
    "Скасування транзакцій (PENDING → CANCELLED) з фіксацією причини",
    "Повернення коштів (COMPLETED → REFUND) з створенням зворотньої транзакції",
    "Перегляд аналітики та звітів по всій системі",
]
for f in admin_funcs:
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent  = Cm(1.25)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p2.add_run(f"— {f}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)


# ════════════════════════════════════════════════════════════
#  4. ЗАПУСК ЗАСТОСУНКУ
# ════════════════════════════════════════════════════════════
add_heading("4. Запуск застосунку")

add_subheading("4.1 Конфігурація")
body("Перед запуском необхідно переконатись у наявності MySQL-сервера "
     "з базою даних course-2026-np. Основні параметри конфігурації у файлі "
     "src/main/resources/application.properties:")

code_lines = [
    "spring.application.name=course-2026-np",
    "server.port=8080",
    "spring.datasource.url=jdbc:mysql://localhost:3306/course-2026-np",
    "spring.datasource.username=root",
    "spring.jpa.hibernate.ddl-auto=update",
    "spring.jpa.show-sql=true",
    "spring.freemarker.suffix=.ftl",
    "spring.freemarker.settings.datetime_format=dd.MM.yyyy HH:mm",
]
for line in code_lines:
    p_code = doc.add_paragraph(style='Normal')
    p_code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_code.paragraph_format.left_indent  = Cm(1.25)
    p_code.paragraph_format.space_before = Pt(0)
    p_code.paragraph_format.space_after  = Pt(0)
    p_code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p_code.add_run(line)
    r.font.name  = 'Courier New'
    r.font.size  = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

add_paragraph("", sb=4)

add_subheading("4.2 Запуск та ініціалізація")
body("Застосунок запускається командою mvn spring-boot:run або через "
     "головний клас Course2026NpApplication у IDE. При першому запуску "
     "Hibernate автоматично створює таблиці у базі даних, "
     "після чого компонент DataInitializer заповнює систему тестовими "
     "даними: адміністратор (admin@payflow.com), менеджер (manager@payflow.com) "
     "та клієнт (client@payflow.com). У консолі відображаються SQL-запити "
     "Hibernate та повідомлення Spring Boot про успішний старт на порту 8080.")

add_screenshot_placeholder("Запуск Spring Boot застосунку в консолі — стартові логи Tomcat")


# ════════════════════════════════════════════════════════════
#  5. ТЕСТУВАННЯ ФУНКЦІОНАЛУ
# ════════════════════════════════════════════════════════════
add_heading("5. Тестування функціоналу")

# ─── 5.1 CLIENT ────────────────────────────────────────────
add_subheading("5.1 Функціонал клієнта (CLIENT)")

body("Для входу як клієнт використовується обліковий запис "
     "client@payflow.com з паролем client123. "
     "Після успішної автентифікації сесія зберігається у HttpSession "
     "з атрибутом userId. Усі маршрути клієнтської зони мають префікс /dashboard/.")

add_paragraph("Вхід в систему:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Форма входу доступна за адресою /login. Після введення коректних "
     "email та пароля (BCrypt-верифікація) система перенаправляє клієнта "
     "на /dashboard. При перевищенні кількості невдалих спроб "
     "LoginAttemptService блокує доступ до облікового запису.")
add_screenshot_placeholder("Форма входу — сторінка /login")

add_paragraph("Управління картками:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("На сторінці /dashboard/cards клієнт бачить перелік своїх карток з "
     "інформацією про номер картки (16 цифр), ім'я власника, термін дії "
     "та поточний статус (активна/заблокована). Кожну картку можна "
     "заблокувати або розблокувати — це змінює поле isActive в сутності CreditCard.")
add_screenshot_placeholder("Список карток клієнта — сторінка /dashboard/cards")

add_paragraph("Зміна PIN-коду:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Зміна PIN реалізована через модальне вікно на сторінці карток. "
     "Запит відправляється як POST /cards/{id}/change-pin з JSON-тілом. "
     "Якщо PIN вже встановлений, поле «Поточний PIN» обов'язкове. "
     "Новий PIN хешується через BCryptPasswordEncoder(10) та зберігається "
     "у полі pinCode сутності CreditCard.")
add_screenshot_placeholder("Модальне вікно зміни PIN-коду картки")

add_paragraph("Переказ коштів:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("На сторінці /dashboard/payment клієнт може здійснити платіж, "
     "поповнення або переказ між рахунками. Операція TRANSFER автоматично "
     "створює дві записи Payment: один для рахунку-відправника (списання) "
     "та один для рахунку-отримувача (зарахування). Перевіряється наявність "
     "достатнього балансу та денний ліміт витрат (spendingLimit).")
add_screenshot_placeholder("Форма здійснення платежу — сторінка /dashboard/payment")

add_paragraph("Перегляд історії транзакцій:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("На сторінці /dashboard/transactions клієнт бачить усі свої транзакції "
     "з інформацією про дату, суму, тип (PAYMENT / REPLENISHMENT / TRANSFER), "
     "статус (PENDING / COMPLETED / FAILED / CANCELLED) та унікальний "
     "ідентифікатор транзакції (transactionId у форматі TXN...).")
add_screenshot_placeholder("Історія транзакцій клієнта — сторінка /dashboard/transactions")

# ─── 5.2 MANAGER ───────────────────────────────────────────
add_subheading("5.2 Функціонал менеджера (MANAGER)")

body("Менеджер автентифікується через /manager/login. Сесія зберігається "
     "з атрибутом managerUser. Менеджер має виключно права на читання — "
     "жодних кнопок для зміни стану транзакцій не відображається.")

add_paragraph("Дашборд зі статистикою:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Головна сторінка менеджера відображає агреговану статистику по "
     "всіх транзакціях системи: загальна кількість, сумарний обсяг, "
     "середнє значення транзакції, розбивка по статусах.")
add_screenshot_placeholder("Дашборд менеджера зі статистикою — /manager/dashboard")

add_paragraph("Моніторинг транзакцій:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("На сторінці /manager/transactions відображається повний список "
     "транзакцій усіх клієнтів системи з можливістю фільтрації та "
     "сортування. Менеджер може переглянути деталі кожної транзакції, "
     "але не може вносити жодних змін.")
add_screenshot_placeholder("Список транзакцій у кабінеті менеджера")

add_paragraph("Аналітика з графіками:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Сторінка аналітики відображає графіки Chart.js із динамікою "
     "транзакцій: розподіл за типами (PAYMENT, REPLENISHMENT, TRANSFER), "
     "обсяг операцій у часовому розрізі.")
add_screenshot_placeholder("Аналітика у кабінеті менеджера з графіками Chart.js")

# ─── 5.3 ADMIN ─────────────────────────────────────────────
add_subheading("5.3 Функціонал адміністратора (ADMIN)")

body("Адміністратор входить через /admin/login. Сесія зберігається "
     "з атрибутом adminUser. Адміністратор має повний доступ до всіх "
     "функцій управління системою.")

add_paragraph("Управління користувачами:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("На сторінці /admin/users адміністратор бачить список усіх "
     "зареєстрованих користувачів з можливістю блокування/розблокування "
     "акаунтів (поле isActive) та зміни ролей (CLIENT / MANAGER / ADMIN). "
     "Усі адміністративні дії логуються з тегом [ADMIN_ACTION] у security.log.")
add_screenshot_placeholder("Список користувачів у панелі адміністратора — /admin/users")

add_paragraph("Аналітика та управління транзакціями:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Адміністратор бачить усі транзакції системи та може виконувати "
     "операції підтвердження, відхилення, скасування та повернення коштів. "
     "Детальний опис цих операцій наведено в розділі 6.")
add_screenshot_placeholder("Панель управління транзакціями адміністратора — /admin/transactions")


# ════════════════════════════════════════════════════════════
#  6. РЕАЛІЗОВАНІ ЗМІНИ
# ════════════════════════════════════════════════════════════
add_heading("6. Реалізовані зміни")

# ─── 6.1 PIN FIX ───────────────────────────────────────────
add_subheading("6.1 Виправлення форми зміни PIN-коду")

body("У початковій версії форма зміни PIN-коду мала проблему: поле "
     "«Поточний PIN» не відображалось умовно залежно від того, чи вже "
     "встановлено PIN на картці. Це призводило до того, що при спробі "
     "встановити PIN вперше форма вимагала введення поточного PIN, якого "
     "ще не існувало.")
body("Виправлення полягало у додаванні умовної логіки у шаблоні FreeMarker "
     "та JavaScript-обробнику модального вікна. При відкритті модального "
     "вікна функція openPinModal(cardId, hasPin) отримує прапорець hasPin "
     "(з поля card.pinCode у шаблоні: '${(card.pinCode?? && card.pinCode?has_content)?c}'). "
     "Поле «Поточний PIN» відображається лише якщо hasPin === true. "
     "На боці сервера (CreditCardController, метод changePin) логіка також "
     "перевіряє наявність PIN та вимагає currentPin тільки у відповідному випадку.")
body("Додатково реалізовано захист: спроба змінити PIN чужої картки "
     "повертає HTTP 403, а невалідований запит без активної сесії — HTTP 401. "
     "Всі успішні та невдалі спроби фіксуються у logах з тегами "
     "[PIN_CHANGE_OK] та [PIN_CHANGE_FAIL].")
add_screenshot_placeholder("Форма зміни PIN-коду після виправлення — коректне відображення полів")

# ─── 6.2 CANCEL ────────────────────────────────────────────
add_subheading("6.2 Скасування транзакцій (Cancellation)")

body("Функціонал скасування транзакцій реалізований у класах "
     "AdminController та PaymentService. Операція доступна виключно "
     "адміністратору через POST-запит на /admin/transactions/cancel.")

add_paragraph("Бізнес-логіка:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Скасувати можна лише транзакції зі статусом PENDING, тобто ті, "
     "що ще очікують на обробку. Це логічне обмеження: транзакції зі "
     "статусом COMPLETED вже виконані та змінили баланс рахунку, тому "
     "для їх відхилення необхідна процедура повернення коштів (refund). "
     "Якщо передана транзакція має статус, відмінний від PENDING, "
     "метод cancelPayment кидає IllegalArgumentException з повідомленням "
     "'Only PENDING transactions can be cancelled'.")

add_paragraph("Статусна модель:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("При скасуванні викликається метод payment.cancel(cancelledBy, cancelReason) "
     "сутності Payment, який змінює статус PENDING → CANCELLED та заповнює "
     "поля: cancelledBy (рядок виду 'admin@payflow.com (ID: 1)'), "
     "cancelledAt (поточний LocalDateTime) та cancelReason (причина скасування, "
     "за замовчуванням 'No reason provided' якщо не вказана).")

add_paragraph("Вплив на баланс:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Оскільки скасовується лише PENDING-транзакція, жодних змін балансу "
     "рахунків не відбувається — кошти ще не були переміщені. "
     "Операція є виключно адміністративним записом у базі даних.")

add_paragraph("Аудит-логування:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Кожне успішне скасування фіксується одночасно у двох лог-файлах "
     "за допомогою виклику log.info та securityLog.info з тегом [CANCEL_OK]. "
     "Формат запису: [CANCEL_OK] paymentId={} amount={} cancelledBy={} reason={}.")

add_screenshot_placeholder("Кнопка Cancel у списку транзакцій адміністратора")
add_screenshot_placeholder("Модальне вікно підтвердження скасування транзакції")
add_screenshot_placeholder("Відображення скасованої транзакції зі статусом CANCELLED у таблиці")

# ─── 6.3 REFUND ────────────────────────────────────────────
add_subheading("6.3 Повернення коштів (Refund)")

body("Функціонал повернення коштів реалізований у класах AdminController "
     "та PaymentService (метод refundPayment). "
     "Операція доступна виключно адміністратору через POST-запит на "
     "/admin/transactions/refund.")

add_paragraph("Бізнес-логіка:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Повернення виконується лише для транзакцій зі статусом COMPLETED. "
     "Додатково перевіряється, чи не було вже здійснено повернення для "
     "цієї транзакції раніше: paymentRepository.existsByOriginalPaymentId(paymentId). "
     "Якщо повернення вже існує — кидається виняток із повідомленням "
     "'Transaction #X has already been refunded'. "
     "Також перевіряється, що рахунок не заблокований (AccountStatus.BLOCKED).")

add_paragraph("Чому оригінальна транзакція не змінюється:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Принцип незмінності фінансового аудит-трейлу: оригінальна транзакція "
     "залишається у стані COMPLETED назавжди. Замість зміни існуючого запису "
     "створюється нова транзакція повернення з полем originalPaymentId, "
     "яке посилається на ID оригінальної транзакції. "
     "Це гарантує, що будь-яка зовнішня система або аудитор завжди може "
     "відстежити повний ланцюжок: що було нараховано, а що — повернуто.")

add_paragraph("Реверсія балансів за типами транзакцій:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Логіка повернення залежить від типу оригінальної транзакції:")
reversal_items = [
    ("PAYMENT (списання)",
     "кошти повертаються на рахунок — баланс збільшується на суму; "
     "створюється нова транзакція типу REPLENISHMENT"),
    ("REPLENISHMENT (поповнення)",
     "кошти знімаються з рахунку (скасування поповнення) — баланс зменшується; "
     "створюється нова транзакція типу PAYMENT; перевіряється достатність балансу"),
    ("TRANSFER (переказ)",
     "найскладніший випадок: кошти повертаються відправнику (баланс +) і "
     "знімаються з отримувача (баланс −); "
     "створюються ДВІ нові транзакції — PAYMENT для отримувача та REPLENISHMENT для відправника; "
     "перевіряється достатність балансу на рахунку отримувача"),
]
for rtype, rdesc in reversal_items:
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent  = Cm(1.25)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p2.add_run(f"• {rtype}: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.bold = True
    r2 = p2.add_run(rdesc)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)

add_paragraph("Повний цикл статусів транзакції:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("PENDING → COMPLETED (після підтвердження адміном) → "
     "при поверненні коштів оригінальна транзакція залишається COMPLETED, "
     "нова транзакція-повернення одразу отримує статус COMPLETED через "
     "виклик refund.complete(). У полі description зворотньої транзакції "
     "зберігається посилання: 'Refund for TXN...' з оригінальним transactionId.")

add_paragraph("Аудит-логування:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=2)
body("Успішне повернення фіксується у security.log з тегом [REFUND_OK]. "
     "Формат запису: [REFUND_OK] originalPaymentId={} refundId={} amount={} "
     "type={} refundedBy={}.")

add_screenshot_placeholder("Кнопка Refund у списку транзакцій адміністратора")
add_screenshot_placeholder("Модальне вікно підтвердження повернення коштів")
add_screenshot_placeholder("Відображення рядка REFUND-транзакції в таблиці транзакцій")


# ════════════════════════════════════════════════════════════
#  7. АНАЛІЗ ЛОГІВ
# ════════════════════════════════════════════════════════════
add_heading("7. Аналіз логів")

body("Система логування налаштована у файлі logback-spring.xml. "
     "Ведуться два окремих лог-файли у каталозі logs/: "
     "payflow.log (загальний, рівень DEBUG+) та "
     "security.log (аудит безпеки, рівень DEBUG+). "
     "Архівування відбувається при досягненні 10 МБ або зміні дати; "
     "зберігаються архіви за 30 днів, максимальний розмір — 300 МБ (payflow) "
     "та 100 МБ (security).")

add_subheading("7.1 Аналіз payflow.log")
body("Основний лог-файл містить записи про всі операції застосунку. "
     "Формат рядка: дата-час (yyyy-MM-dd HH:mm:ss.SSS), рівень логування, "
     "ім'я потоку та скорочений шлях до класу. "
     "Записи класу ua.com.kisit ведуться з рівнем DEBUG, що дозволяє "
     "бачити SQL-запити Hibernate та деталі бізнес-операцій.")

add_subheading("7.2 Аналіз security.log")
body("Файл безпеки містить аудит усіх критичних подій системи з чіткими "
     "тегами у форматі [EVENT_TYPE]. Кожна запис у security.log доповнюється "
     "префіксом [SECURITY], що дозволяє легко виокремити записи безпеки "
     "з агрегованих логів. Типові теги для операцій, реалізованих у роботі:")

security_events = [
    "[LOGIN_OK] / [LOGIN_FAIL]",
    "[ADMIN_LOGIN_OK] / [MANAGER_LOGIN_OK]",
    "[LOGOUT]",
    "[BRUTE_FORCE_BLOCK]",
    "[UNAUTHENTICATED_ACCESS] — доступ без активної сесії (HTTP 401)",
    "[ACCESS_DENIED_403] — спроба доступу до чужих ресурсів",
    "[PIN_CHANGE_OK] / [PIN_CHANGE_FAIL]",
    "[CANCEL_OK] — успішне скасування PENDING-транзакції",
    "[REFUND_OK] — успішне повернення коштів по COMPLETED-транзакції",
    "[ADMIN_ACTION] — блокування/розблокування користувача, зміна ролі",
    "[CARD_TOGGLE_BLOCK] — блокування/розблокування картки",
]
for evt in security_events:
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent  = Cm(1.25)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p2.add_run(f"— {evt}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

add_screenshot_placeholder("Вміст security.log з записами [CANCEL_OK] та [REFUND_OK]")


# ════════════════════════════════════════════════════════════
#  8. ВИСНОВОК
# ════════════════════════════════════════════════════════════
add_heading("8. Висновок")

body("У ході виконання лабораторної роботи №20 було успішно виконано "
     "запуск, функціональне тестування та аналіз логів веб-застосунку "
     "PayFlow — навчальної банківської платіжної системи на Spring Boot 4.0.2.")
body("В результаті роботи були досліджені та практично перевірені такі аспекти:")

conclusions = [
    ("Структура Spring Boot застосунку",
     "трирівнева архітектура Controller → Service → Repository, "
     "автоматична конфігурація, lifecycle від старту до готовності до запитів"),
    ("Система автентифікації та авторизації",
     "сесійна автентифікація через HttpSession, три ролі з різними рівнями "
     "доступу, захист від brute-force атак, BCrypt-хешування паролів та PIN"),
    ("Управління фінансовими транзакціями",
     "повний цикл транзакції PENDING → COMPLETED, три типи операцій "
     "(PAYMENT, REPLENISHMENT, TRANSFER), перевірка балансу та лімітів"),
    ("Реалізація скасування транзакцій (Cancel)",
     "операція можлива лише для PENDING-транзакцій, "
     "фіксація причини та виконавця у полях cancelledBy, cancelledAt, cancelReason"),
    ("Реалізація повернення коштів (Refund)",
     "операція можлива лише для COMPLETED-транзакцій, "
     "незмінність оригінальних записів як основа аудит-трейлу, "
     "створення зворотніх транзакцій з originalPaymentId для повного ланцюжку"),
    ("Важливість аудит-трейлу у фінансових системах",
     "принцип незмінності фінансових записів, "
     "відокремлений security.log для аудиту, теговані записи [EVENT_TYPE] "
     "для швидкого пошуку подій безпеки та адміністративних дій"),
]
for title, desc in conclusions:
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent  = Cm(0)
    p2.paragraph_format.first_line_indent = Cm(1.25)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p2.add_run(f"{title}: ")
    r1.font.name  = 'Times New Roman'
    r1.font.size  = Pt(14)
    r1.font.bold  = True
    r2 = p2.add_run(desc + ".")
    r2.font.name  = 'Times New Roman'
    r2.font.size  = Pt(14)

body("Отримані знання про організацію фінансових операцій у Spring Boot, "
     "принципи аудит-трейлу та різницю між скасуванням і поверненням "
     "транзакцій є практично застосовними при розробці реальних "
     "банківських та платіжних систем.")


# ════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════
import os
out_dir = r"C:\Users\Lenovo\OneDrive\Desktop\practica"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "ЛР_20_Зівак_Сергій_371.docx")
doc.save(out_path)
print(f"Збережено: {out_path}")
print(f"Кількість рисунків (заглушок): {fig_counter[0]}")
